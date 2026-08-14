"""
报告生成系统 —— 分层知识库 + 来源追溯 + 自检验证
======================================================
每一条结论都标注依据来源，生成完后自动验证是否可追溯。
"""
import os, sys, re, time, json, hashlib
from typing import List, Dict, Tuple, Optional

if sys.platform == "win32":
    try: sys.stdout.reconfigure(encoding="utf-8"); sys.stderr.reconfigure(encoding="utf-8")
    except: pass

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 图表
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

# 中文字体
for fname in ["SimHei", "Microsoft YaHei", "Noto Sans CJK SC", "DejaVu Sans"]:
    try: plt.rcParams["font.sans-serif"] = [fname]; break
    except: pass
plt.rcParams["axes.unicode_minus"] = False

# ---- 知识库分层加载 ----
class LayeredKnowledge:
    """
    三层知识库：
      L1 政策法规: 货币政策报告、监管文件 → 用于宏观背景、政策解读
      L2 专业标准: 风险指标调研、RAG架构调研 → 用于风险定义、方法论
      L3 实例数据: 股票风险 JSON → 用于具体股票分析
    """

    def __init__(self, folder: str = "./knowledge"):
        self.folder = folder
        self.l1_policy = []      # [(filename, content)]
        self.l2_standards = []   # [(filename, content)]
        self.l3_cases = []       # [(filename, data_dict)] for JSON
        self.all_texts = []      # [(filename, content)] for retrieval
        self._load()

    def _load(self):
        if not os.path.exists(self.folder):
            return
        for fn in os.listdir(self.folder):
            fp = os.path.join(self.folder, fn)
            try:
                if fn.endswith(".txt"):
                    with open(fp, "r", encoding="utf-8") as f:
                        content = f.read()
                    if any(kw in fn for kw in ["货币", "季度", "政策", "监管", "合规"]):
                        self.l1_policy.append((fn, content))
                    else:
                        self.l2_standards.append((fn, content))
                    self.all_texts.append((fn, content))
                elif fn.endswith(".md"):
                    with open(fp, "r", encoding="utf-8") as f:
                        content = f.read()
                    self.l2_standards.append((fn, content))
                    self.all_texts.append((fn, content))
                elif fn.endswith(".json"):
                    with open(fp, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if isinstance(data, dict) and "stock" in data:
                        self.l3_cases.append((fn, data))
                        # 转自然语言文本供检索
                        self.all_texts.append((fn, self._json_to_text(fn, data)))
                    elif isinstance(data, list):
                        for item in data:
                            if isinstance(item, dict):
                                self.l3_cases.append((fn, item))
                                self.all_texts.append((fn, self._json_to_text(fn, item)))
                elif fn.endswith(".docx"):
                    try:
                        doc = Document(fp)
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        self.l2_standards.append((fn, text))
                        self.all_texts.append((fn, text))
                    except:
                        pass
            except Exception as e:
                print(f"  [Skip] {fn}: {e}")

    @staticmethod
    def _json_to_text(filename: str, data: dict) -> str:
        stock = data.get("stock", "?")
        summary = data.get("rag_summary", "")
        summary = re.sub(r"^rag[_\s]?summary[：:]?\s*", "", summary, flags=re.IGNORECASE)
        return (
            f"股票{stock}。"
            f"风险等级{data.get('final_level','')}。"
            f"处置建议：{data.get('action','')}。"
            f"主导风险来源：{data.get('dominant_source','')}。"
            f"AI评分{data.get('ai_score','')}分。"
            f"波动率{data.get('hist_vol_pct','')}%。"
            f"最大回撤{data.get('max_drawdown_pct','')}%。"
            f"负债率{data.get('debt_ratio_pct','')}%。"
            f"流动比率{data.get('current_ratio','')}。"
            f"复合评分{data.get('composite_score','')}。"
            f"详情：{summary}"
        )

    def summary(self) -> str:
        return (f"L1政策: {len(self.l1_policy)} | "
                f"L2标准: {len(self.l2_standards)} | "
                f"L3实例: {len(self.l3_cases)} | "
                f"文本总量: {len(self.all_texts)}")


# ---- 简易检索器（纯 TF-IDF，不依赖额外模型）----
import jieba
import jieba.analyse
from sklearn.feature_extraction.text import TfidfVectorizer

def _extract_kw(text, topK=3):
    try:
        return jieba.analyse.extract_tags(text, topK=topK)
    except:
        return [w for w in text[:30].replace(" ", "") if len(w) > 1]
from sklearn.metrics.pairwise import cosine_similarity

class SimpleRetriever:
    def __init__(self, docs: List[Tuple[str, str]]):
        self.docs = docs
        self.texts = [d[1] for d in docs]
        self.sources = [d[0] for d in docs]
        if self.texts:
            self.vec = TfidfVectorizer(max_features=2000, tokenizer=lambda x: jieba.lcut(x))
            self.matrix = self.vec.fit_transform(self.texts)
        else:
            self.vec = None
            self.matrix = None

    def search(self, query: str, top_k: int = 5) -> List[Dict]:
        if self.vec is None:
            return []
        qv = self.vec.transform([query])
        scores = cosine_similarity(qv, self.matrix).flatten()
        idxs = scores.argsort()[-top_k:][::-1]
        return [{"content": self.texts[i], "source": self.sources[i], "score": float(scores[i])}
                for i in idxs if scores[i] > 0.01]


# ---- 图表生成 ----
class ChartMaker:
    """从 JSON 数据生成图表 PNG"""
    def __init__(self, out_dir="./charts"):
        self.out_dir = out_dir
        os.makedirs(out_dir, exist_ok=True)

    def _ensure_stock_dir(self, stock_code: str = None) -> str:
            """获取/创建股票子目录"""
            if stock_code:
                sub_dir = os.path.join(self.out_dir, stock_code)
                os.makedirs(sub_dir, exist_ok=True)
                return sub_dir
            return self.out_dir
    

    def all(self, stock: Dict, prefix: str = "chart") -> Dict[str, str]:
        """生成全部图表，返回 {名称: 路径}"""
        stock_code = stock.get("stock")
        charts = {}
        p = self._radar(stock, f"{prefix}_radar", stock_code)
        if p: charts["radar"] = p
        p = self._bar(stock, f"{prefix}_bar", stock_code)
        if p: charts["bar"] = p
        p = self._mc(stock, f"{prefix}_mc", stock_code)
        if p: charts["mc"] = p
        return charts

    def _radar(self, d: Dict, name: str,stock_code: str = None) -> Optional[str]:
        """风险维度雷达图 —— 展示各维度等级"""
        dims = ["波动率风险", "回撤风险", "财务杠杆", "流动性", "增长"]
        levels = {"Ⅰ级": 1, "Ⅱ级": 2, "Ⅲ级": 3, "Ⅳ级": 4}
        vals = [
            levels.get(d.get("vol_level", ""), 2),
            levels.get(d.get("mdd_level", ""), 2),
            levels.get(d.get("fin_level_debt", ""), 2),
            levels.get(d.get("fin_level_liquidity", ""), 2),
            levels.get(d.get("fin_level_growth", ""), 2),
        ]
        N = len(dims)
        angles = [n / N * 2 * np.pi for n in range(N)]
        angles += angles[:1]
        vals += vals[:1]

        fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
        ax.fill(angles, vals, alpha=0.25, color="#c5221f")
        ax.plot(angles, vals, "o-", linewidth=2, color="#c5221f")
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(dims, fontsize=10)
        ax.set_yticks([1, 2, 3, 4])
        ax.set_yticklabels(["Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ"], fontsize=8)
        ax.set_ylim(0, 4)
        ax.set_title(f'{d.get("stock","?")} 风险雷达图', fontsize=13, fontweight="bold", pad=20)
        for a, v in zip(angles[:-1], vals[:-1]):
            ax.annotate(f"L{v}", (a, v+0.15), ha="center", fontsize=9)
        plt.tight_layout()
        return self._save(fig, name,stock_code)

    def _bar(self, d: Dict, name: str,stock_code: str = None) -> Optional[str]:
        """关键指标 vs 阈值柱状图"""
        metrics = ["波动率%", "最大回撤%", "负债率%"]
        values = [
            float(d.get("hist_vol_pct", 0)),
            float(d.get("max_drawdown_pct", 0)),
            float(d.get("debt_ratio_pct", 0)),
        ]
        thresholds = [25, 20, 80]  # 极端阈值

        fig, ax = plt.subplots(figsize=(7, 4))
        x = np.arange(len(metrics))
        w = 0.35
        bars1 = ax.bar(x - w/2, values, w, label="当前值", color="#c5221f")
        bars2 = ax.bar(x + w/2, thresholds, w, label="阈值", color="#aaaaaa", alpha=0.5)
        for bar, val in zip(bars1, values):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+1, f"{val:.1f}", ha="center", fontsize=10)
        ax.set_xticks(x)
        ax.set_xticklabels(metrics, fontsize=10)
        ax.set_title(f'{d.get("stock","?")} 关键指标 vs 阈值', fontsize=13, fontweight="bold")
        ax.legend(fontsize=9)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        ax.yaxis.grid(True, alpha=0.3)
        plt.tight_layout()
        return self._save(fig, name, stock_code)

    def _mc(self, d: Dict, name: str,stock_code: str = None) -> Optional[str]:
        """蒙特卡洛风险分布 —— 柱状图"""
        labels = ["1%最坏", "5%最坏", "10%最坏"]
        vals = [
            abs(float(d.get("mc_q1_1d_pct", 0) or 0)),
            abs(float(d.get("mc_q5_1d_pct", 0) or 0)),
            abs(float(d.get("mc_q10_1d_pct", 0) or 0)),
        ]
        if sum(vals) == 0:
            return None
        fig, ax = plt.subplots(figsize=(5, 3.5))
        colors = ["#c5221f", "#e37400", "#f9ab00"]
        bars = ax.bar(labels, vals, color=colors, edgecolor="white")
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.1, f"{val:.2f}%", ha="center", fontsize=11)
        ax.set_title(f'{d.get("stock","?")} 蒙特卡洛日亏损分布', fontsize=12, fontweight="bold")
        ax.set_ylabel("日亏损 %", fontsize=10)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        plt.tight_layout()
        return self._save(fig, name,stock_code)
    
    def _save(self, fig, name: str, stock_code: str = None) -> Optional[str]:
        """保存图片到指定目录（支持股票子目录）"""
        save_dir = self._ensure_stock_dir(stock_code)
        path = os.path.join(save_dir, f"{name}.png")
        try:
            fig.savefig(path, dpi=120, bbox_inches="tight", facecolor="white")
            plt.close(fig)
            return path
        except:
            plt.close(fig)
            return None


# ---- 报告生成器 ----
class ReportGenerator:
    """
    报告结构（自上而下定义）：
      一、基本信息        ← L3 JSON 结构化数据
      二、风险评估        ← L3 数据 + L2 标准
      三、数据解读        ← L3 数值 + L2 方法论
      四、处置建议        ← L3 action 字段 + L1 政策依据
      五、依据溯源        ← 上述所有信息来源的汇总
      六、合规声明
    """

    REPORT_SECTIONS = [
        {"id": "basic_info",    "title": "基本信息",    "depends": ["L3"], "need_table": True},
        {"id": "risk_assess",   "title": "风险评估",    "depends": ["L3", "L2"], "need_table": True},
        {"id": "data_reading",  "title": "数据解读",    "depends": ["L3", "L2", "L1"], "need_table": True},
        {"id": "actions",       "title": "处置建议",    "depends": ["L3", "L1"], "need_table": False},
        {"id": "sources",       "title": "依据溯源",    "depends": ["L1", "L2", "L3"], "need_table": True},
        {"id": "disclaimer",    "title": "合规声明",    "depends": [], "need_table": False},
    ]

    def __init__(self, kb_folder: str = "./knowledge", risk_data: dict = None):
        self.kb = LayeredKnowledge(kb_folder)
        self.risk_data = risk_data  # 保存原始 risk_data
        self.retriever = SimpleRetriever(self.kb.all_texts)
        print(f"  [KB] {self.kb.summary()}")

    def generate(self, topic: str) -> str:
        t0 = time.time()
        topic_clean = topic.strip()

        # ========== 【核心修改】优先使用传入的 risk_data ==========
        target_stock = None
    
        # 1. 如果存在传入的 risk_data，直接作为目标数据（无需读文件）
        if self.risk_data and self.risk_data.get('stock'):
            target_stock = self.risk_data
            print(f"DEBUG: 使用传入的 risk_data 作为 target_stock")
            print(f"DEBUG: sector_avg_vol_pct = {target_stock.get('sector_avg_vol_pct')}")
        else:
            # 2. 如果 risk_data 没有，再从 l3_cases 匹配
            for fn, data in self.kb.l3_cases:
                code = data.get("stock", "")
                if code and code in topic_clean:
                    target_stock = data
                    break

        if target_stock:
            print(f"DEBUG: target_stock keys: {list(target_stock.keys())}")
            print(f"DEBUG: sector_avg_vol_pct = {target_stock.get('sector_avg_vol_pct')}")

        # ---- 图表生成 ----
        chart_maker = ChartMaker()
        charts = {}
        if target_stock:
            charts = chart_maker.all(target_stock, f"stock_{target_stock.get('stock','x')}")

        # ---- 逐节生成 ----
        sections = []
        all_sources = set()

        for sec in self.REPORT_SECTIONS:
            content, sources, table = self._build_section(sec, topic_clean, target_stock)
            # 图表分配
            chart = None
            sid = sec["id"]
            if sid == "risk_assess" and "radar" in charts:
                chart = charts["radar"]
            elif sid == "data_reading":
                chart = charts.get("bar") or charts.get("mc")

            sections.append({
                "title": sec["title"],
                "content": content,
                "sources": sources,
                "table": table,
                "chart": chart,
            })
            all_sources.update(sources)

        # ---- 自检验证 ----
        verify_result = self._verify(sections)

        # ---- 导出 Word ----
        os.makedirs("./reports", exist_ok=True)
        ts = time.strftime("%Y%m%d_%H%M%S")
        docx_path = f"./reports/report_{ts}.docx"
        self._export_docx(docx_path, topic_clean, sections, verify_result, all_sources,
                          target_stock, (time.time()-t0)*1000)

        # 同时存 txt
        txt_path = docx_path.replace(".docx", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(self._format_txt(topic_clean, sections, verify_result, all_sources))

        elapsed = (time.time() - t0) * 1000
        print(f"\n  Report: {docx_path}  ({elapsed:.0f}ms)")
        return docx_path

    def _build_section(self, sec: Dict, topic: str, stock: Optional[Dict]) -> Tuple[str, List[str], Optional[Tuple]]:
        """构建一个章节，有结构化数据优先用数据，否则用RAG检索"""
        sid = sec["id"]
        lines = []
        sources = []
        table = None
        has_stock = stock is not None

        # ---- 无股票数据时的通用处理 ----
        if not has_stock and sid not in ("sources", "disclaimer"):
            rag = self.retriever.search(topic, top_k=5)
            if rag:
                lines.append(f"知识库检索结果（主题「{topic}」相关）：")
                for r in rag[:5]:
                    excerpt = r["content"][:250].replace("\n", " ")
                    lines.append(f"  · {excerpt}")
                    lines.append(f"    ——来源: {r['source']}")
                    sources.append(r["source"])
                lines.append("")
                lines.append(f"【说明】未找到与「{topic}」精确匹配的结构化数据，以上为知识库语义检索结果。")
            else:
                lines.append(f"知识库中未找到与「{topic}」相关的信息。")
                lines.append("请尝试: 1) 输入具体股票代码如000333 2) 输入更具体的金融主题")
            return "\n".join(lines), sources, None

        if sid == "basic_info" and has_stock:
            lines.append(f"分析对象：股票 {stock.get('stock','?')}")
            lines.append(f"评估日期：{stock.get('eval_date','?')}")
            lines.append(f"风险等级：{stock.get('final_level','?')} (数值{stock.get('final_level_num','?')})")
            lines.append(f"AI综合评分：{stock.get('ai_score','?')}分 (对应{stock.get('ai_level','?')})")
            lines.append(f"复合评分：{stock.get('composite_score','?')}")
            lines.append(f"主导风险来源：{stock.get('dominant_source','?')}")
            # ========== 【新增】置信度展示 ==========
            conf = stock.get('confidence_score')
            if conf is not None and not pd.isna(conf):
                conf_pct = conf * 100
                if conf < 0.35:
                    tip = "⚠️ 当前市场环境与训练数据差异巨大，AI预测置信度极低，请以实时波动率与财务指标为准！"
                elif conf < 0.65:
                    tip = "⚡ AI模型置信度偏低，建议结合实时数据进行交叉验证。"
                else:
                    tip = "✅ 当前模型置信度较高，预测结果可作为有效参考。"
                lines.append(f"AI模型置信度：{conf_pct:.1f}%（{tip}）")
            else:
                lines.append(f"AI模型置信度：N/A（无法评估）")

            lines.append("")
            lines.append("【来源依据】L3-股票风险实例数据")
            sources.append("L3-股票风险JSON")
            # 数据表
            table = (
                ["指标", "数值", "说明"],
                [
                    ["最终风险等级", str(stock.get('final_level','')), "综合评估结果"],
                    ["AI评分", f"{stock.get('ai_score','')}分", "LSTM模型综合评分"],
                    ["波动率等级", str(stock.get('vol_level','')), "历史年化波动率分级"],
                    ["回撤等级", str(stock.get('mdd_level','')), "历史最大回撤分级"],
                    ["负债等级", str(stock.get('fin_level_debt','')), "资产负债率分级"],
                    ["流动性等级", str(stock.get('fin_level_liquidity','')), "流动比率分级"],
                    ["增长等级", str(stock.get('fin_level_growth','')), "连续负增长分级"],
                ]
            )

        elif sid == "risk_assess":
            if not has_stock:
                return self._fallback_section(sid, topic, sources)
            # 从 L3 取风险数据 + L2 取标准定义
                lines.append(f"根据风险评估模型，该股票被评定为「{stock.get('final_level','')}」。")
                lines.append(f"处置动作：{stock.get('action','')}")
                lines.append(f"裁决状态：{stock.get('verdict_status','')}")
                lines.append(f"冲突状态：{'存在冲突' if stock.get('has_conflict') else '无冲突'}"
                           f" (冲突数: {stock.get('conflict_count',0)})")
                lines.append("")
                lines.append("【来源依据】L3-股票风险实例数据")
                sources.append("L3-股票风险JSON")

            # L2 风险标准
            std_results = self.retriever.search("风险等级 分级标准 阈值 指标定义", top_k=3)
            if std_results:
                lines.append("风险分级标准参考：")
                for r in std_results[:2]:
                    excerpt = r["content"][:200].replace("\n", " ")
                    lines.append(f"  · {excerpt}")
                    lines.append(f"    (来源: {r['source']})")
                    sources.append(r["source"])
                lines.append("")

            table = (
                ["风险维度", "等级", "关键指标", "依据来源"],
                [
                    ["波动率风险", str(stock.get('vol_level','')), f"{stock.get('hist_vol_pct','')}%", "L3-历史数据"],
                    ["回撤风险", str(stock.get('mdd_level','')), f"{stock.get('max_drawdown_pct','')}%", "L3-历史数据"],
                    ["财务杠杆风险", str(stock.get('fin_level_debt','')), f"{stock.get('debt_ratio_pct','')}%", "L3-财务数据"],
                    ["流动性风险", str(stock.get('fin_level_liquidity','')), f"{stock.get('current_ratio','')}", "L3-财务数据"],
                    ["增长风险", str(stock.get('fin_level_growth','')), f"{stock.get('consecutive_neg_quarters','')}季负增长", "L3-财务数据"],
                ]
            ) if stock else None

        elif sid == "data_reading":
            if stock:
                print(f"DEBUG: data_reading - stock keys: {list(stock.keys())}")
                print(f"DEBUG: sector_avg_vol_pct = {stock.get('sector_avg_vol_pct')}")

                lines.append(f"波动率分析：历史年化波动率 {stock.get('hist_vol_pct','')}%，"
                           f"蒙特卡洛预测最坏日亏损 {stock.get('mc_q1_1d_pct','')}%。")
                lines.append(f"回撤分析：历史最大回撤 {stock.get('max_drawdown_pct','')}%，"
                           f"动态阈值 {stock.get('dynamic_threshold_pct','')}%。")
                lines.append(f"财务分析：资产负债率 {stock.get('debt_ratio_pct','')}%，"
                           f"流动比率 {stock.get('current_ratio','')}。")
                lines.append("")
                lines.append("【来源依据】L3-股票风险JSON中的量化数据字段")
                sources.append("L3-股票风险JSON")

                # ========== 【新增】行业对比信息 ==========
                sector_avg_vol = stock.get('sector_avg_vol_pct')
                sector_avg_mdd = stock.get('sector_avg_mdd_pct')
                sector = stock.get('sector', '未知行业')
        
                if sector_avg_vol is not None and not pd.isna(sector_avg_vol):
                    diff_vol = stock.get('hist_vol_pct', 0) - sector_avg_vol
                    if diff_vol > 0:
                        vol_compare = f"高于行业均值（{sector}平均 {sector_avg_vol:.1f}%，高出 {diff_vol:.1f} 个百分点）"
                    else:
                        vol_compare = f"低于行业均值（{sector}平均 {sector_avg_vol:.1f}%，低 {-diff_vol:.1f} 个百分点）"
                    lines.append(f"行业对比：波动率 {vol_compare}。")
        
                if sector_avg_mdd is not None and not pd.isna(sector_avg_mdd):
                    diff_mdd = stock.get('max_drawdown_pct', 0) - sector_avg_mdd
                    if diff_mdd > 0:
                        mdd_compare = f"高于行业均值（{sector}平均 {sector_avg_mdd:.1f}%，高出 {diff_mdd:.1f} 个百分点）"
                    else:
                        mdd_compare = f"低于行业均值（{sector}平均 {sector_avg_mdd:.1f}%，低 {-diff_mdd:.1f} 个百分点）"
                    lines.append(f"行业对比：回撤 {mdd_compare}。")
        
                lines.append("")
                lines.append("【来源依据】L3-股票风险JSON中的量化数据字段 + 行业平均水平")
                sources.append("L3-股票风险JSON")
                sources.append("行业平均水平")

            # L2 方法论
            method_results = self.retriever.search("波动率 回撤 资产负债率 评估方法 解释", top_k=3)
            if method_results:
                lines.append("指标解读参考：")
                for r in method_results[:2]:
                    excerpt = r["content"][:200].replace("\n", " ")
                    lines.append(f"  · {excerpt}")
                    lines.append(f"    (来源: {r['source']})")
                    sources.append(r["source"])
                lines.append("")

            # 在 data_reading 表格中
            table = (
            ["指标", "当前值", "等级", "行业均值", "来源"],
            [
                ["波动率", f"{stock.get('hist_vol_pct','')}%", str(stock.get('vol_level','')),
                f"{stock.get('sector_avg_vol_pct', '-')}%" if stock.get('sector_avg_vol_pct') is not None and not pd.isna(stock.get('sector_avg_vol_pct')) else "-",
                "L3"],
                ["最大回撤", f"{stock.get('max_drawdown_pct','')}%", str(stock.get('mdd_level','')),
                f"{stock.get('sector_avg_mdd_pct', '-')}%" if stock.get('sector_avg_mdd_pct') is not None and not pd.isna(stock.get('sector_avg_mdd_pct')) else "-",
                "L3"],
                ["资产负债率", f"{stock.get('debt_ratio_pct','')}%", str(stock.get('fin_level_debt','')),
                f"{stock.get('sector_avg_debt_pct', '-')}%" if stock.get('sector_avg_debt_pct') is not None and not pd.isna(stock.get('sector_avg_debt_pct')) else "-",
                "L3"],
                ["流动比率", str(stock.get('current_ratio','')), str(stock.get('fin_level_liquidity','')),
                str(stock.get('sector_avg_current_ratio', '-')) if stock.get('sector_avg_current_ratio') is not None and not pd.isna(stock.get('sector_avg_current_ratio')) else "-",
                "L3"],
            ]
        ) if stock else None

        elif sid == "actions":
            if stock:
                lines.append(f"基于风险评估结果（{stock.get('final_level','')}），建议采取以下措施：")
                lines.append("")
                lines.append(f"主要处置动作：{stock.get('action','')}")
                lines.append(f"裁决状态：{stock.get('verdict_status','')}")
                lines.append("")
                lines.append("【来源依据】L3-风险系统的执行层输出")
                sources.append("L3-股票风险JSON")

            # L1 政策依据
            policy_results = self.retriever.search("金融风险 防范化解 监管 政策 措施", top_k=3)
            if policy_results:
                lines.append("相关政策依据：")
                for r in policy_results[:2]:
                    excerpt = r["content"][:200].replace("\n", " ")
                    lines.append(f"  · {excerpt}")
                    lines.append(f"    (来源: {r['source']})")
                    sources.append(r["source"])
                lines.append("")

        elif sid == "sources":
            lines.append("本报告各章节依据以下知识库层次：")
            lines.append("")
            lines.append("L1 政策法规层：")
            for fn, _ in self.kb.l1_policy:
                lines.append(f"  - {fn}")
            lines.append("L2 专业标准层：")
            for fn, _ in self.kb.l2_standards:
                lines.append(f"  - {fn}")
            lines.append("L3 实例数据层：")
            for fn, _ in self.kb.l3_cases:
                lines.append(f"  - {fn}")
            lines.append("")
            lines.append("【说明】每条分析结论均标注了对应的依据层和来源文件。")
            table = (
                ["知识层", "文件", "用途"],
                [(f"L1", fn, "宏观政策背景") for fn, _ in self.kb.l1_policy] +
                [(f"L2", fn, "风险标准与方法论") for fn, _ in self.kb.l2_standards] +
                [(f"L3", fn, "股票实例数据") for fn, _ in self.kb.l3_cases]
            )

        elif sid == "disclaimer":
            lines.append("本报告基于以下知识库自动生成：")
            lines.append(f"  L1 政策法规层：{len(self.kb.l1_policy)} 份文档")
            lines.append(f"  L2 专业标准层：{len(self.kb.l2_standards)} 份文档")
            lines.append(f"  L3 实例数据层：{len(self.kb.l3_cases)} 条记录")
            lines.append("")
            lines.append("所有分析结论均标注了依据来源。")
            lines.append("本报告不构成投资建议，使用者应独立核实关键数据。")
            lines.append("报告生成方不对因使用本报告而产生的任何损失承担责任。")

        return "\n".join(lines), list(set(sources)), table

    def _verify(self, sections: List[Dict]) -> List[str]:
        """自检：逐条验证结论是否能在知识库中找到原文依据"""
        result = []
        # 构建全量知识库文本用于检索
        all_kb_text = "\n".join(t for _, t in self.kb.all_texts)

        for sec in sections:
            if sec["title"] in ("合规声明", "依据溯源"):
                continue

            content = sec.get("content", "")
            # 提取所有实质性结论句（以句号、换行分隔，排除【】标记行）
            claims = []
            for line in content.split("\n"):
                line = line.strip()
                if not line or line.startswith("【") or line.startswith("（"):
                    continue
                for sent in re.split(r"[。；]", line):
                    sent = sent.strip()
                    if len(sent) >= 15:
                        claims.append(sent)

            verified = 0
            unchecked = 0
            unverified_claims = []

            for claim in claims:
                # 跳过纯数据展示行和来源标注行
                if re.match(r"^[-\s]*[·•\-]", claim):
                    continue
                if re.match(r"^\d{4}年|^[a-zA-Z]+:|^分析对象|^评估日期|^风险等级|^处置动作|^裁决状态|^冲突状态|^AI综合评分|^复合评分|^主导|^波动率分析|^回撤分析|^财务分析|^主要处置|^基于风险|^\(", claim):
                    continue
                if "来源" in claim[:10] or "source" in claim[:10].lower():
                    continue

                # 宽松匹配：关键词有一个命中即算有依据
                kws = _extract_kw(claim, topK=3)
                found = any(kw in all_kb_text for kw in kws if len(kw) >= 2)

                if found:
                    verified += 1
                else:
                    unverified_claims.append(claim[:60])
                    unchecked += 1

            total = verified + unchecked
            if total == 0:
                result.append(f"✓ {sec['title']}: 全部来自结构化数据，天然有依据")
            elif unchecked == 0:
                result.append(f"✓ {sec['title']}: {verified}/{total} 条结论已在知识库中找到依据")
            elif unchecked <= total * 0.3:
                result.append(f"△ {sec['title']}: {verified}/{total} 有依据, {unchecked} 条未确认"
                           f"{' (' + '; '.join(unverified_claims[:2]) + '...)' if unverified_claims else ''}")
            else:
                result.append(f"✗ {sec['title']}: 仅 {verified}/{total} 有依据, {unchecked} 条可疑"
                           f"{' (' + '; '.join(unverified_claims[:2]) + '...)' if unverified_claims else ''}")

        return result

    def _export_docx(self, path: str, topic: str, sections: List[Dict],
                     verify: List[str], sources: set, stock: Optional[Dict], ms: float):
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

        # 封面
        for _ in range(6): doc.add_paragraph("")
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("数据分析与风险评估报告"); r.bold = True; r.font.size = Pt(24)
        s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run(f"主题: {topic}"); r.font.size = Pt(14)
        s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s2.add_run(f"日期: {time.strftime('%Y-%m-%d %H:%M')}  |  用时: {ms:.0f}ms"); r.font.size = Pt(10)
        doc.add_page_break()

        # 各章节
        for sec in sections:
            h = doc.add_heading(sec["title"], level=1)
            for run in h.runs: run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6d)

            for para in sec["content"].split("\n"):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.line_spacing = 1.5

            # 嵌入图表
            if sec.get("chart") and os.path.exists(sec["chart"]):
                doc.add_paragraph("")
                try:
                    doc.add_picture(sec["chart"], width=Inches(4.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except: pass

            if sec.get("table"):
                headers, rows = sec["table"]
                if rows:
                    doc.add_paragraph("")
                    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
                    tbl.style = "Light Grid Accent 1"
                    for j, h in enumerate(headers): tbl.rows[0].cells[j].text = h
                    for i, row in enumerate(rows):
                        for j, val in enumerate(row):
                            tbl.rows[i+1].cells[j].text = str(val)[:100]

            if sec["title"] != "合规声明":
                doc.add_paragraph("")

        # 自检结果
        doc.add_page_break()
        h = doc.add_heading("附：自检验证报告", level=1)
        for run in h.runs: run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6d)
        for v in verify:
            doc.add_paragraph(v)

        # 页脚
        footer = section.footer
        footer.paragraphs[0].text = f"RAG Report Generator | {len(self.kb.all_texts)} sources"
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(path)

    def _format_txt(self, topic: str, sections: List[Dict], verify: List[str],
                    sources: set) -> str:
        lines = []
        lines.append("=" * 60)
        lines.append(f"  数据分析与风险评估报告")
        lines.append(f"  主题: {topic}")
        lines.append(f"  日期: {time.strftime('%Y-%m-%d %H:%M')}")
        lines.append("=" * 60)
        for sec in sections:
            lines.append(f"\n{'─' * 50}")
            lines.append(f"  {sec['title']}")
            lines.append(f"{'─' * 50}")
            lines.append(sec["content"])
            if sec.get("table"):
                headers, rows = sec["table"]
                if rows:
                    lines.append(f"\n  | {' | '.join(headers)} |")
                    for row in rows:
                        lines.append(f"  | {' | '.join(str(c)[:30] for c in row)} |")
        lines.append(f"\n{'─' * 50}")
        lines.append("  附：自检验证")
        lines.append(f"{'─' * 50}")
        for v in verify:
            lines.append(f"  {v}")
        return "\n".join(lines)


# ============================================================
# 【新增】API 调用入口（兼容 report_api.py）
# ============================================================
def generate_report_from_topic(topic: str, risk_data: dict = None,
                                knowledge_dir: str = "./knowledge") -> dict:
    """
    供 report_api.py 调用的报告生成函数
    """
    # 直接传递 risk_data 给 ReportGenerator
    gen = ReportGenerator(knowledge_dir, risk_data=risk_data)
    docx_path = gen.generate(topic)

    txt_path = docx_path.replace(".docx", ".txt")

    chart_files = []
    charts_dir = "./charts"
    if os.path.exists(charts_dir):
        for root, dirs, files in os.walk(charts_dir):
            for f in files:
                if f.endswith('.png'):
                    rel_path = os.path.relpath(os.path.join(root, f), start=".")
                    chart_files.append(rel_path.replace("\\", "/"))

    preview_text = ""
    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as f:
            preview_text = f.read()

    return {
        'docx_path': docx_path,
        'txt_path': txt_path,
        'chart_files': chart_files,
        'preview_text': preview_text
    }

# ---- CLI ----
def main():
    print("=" * 50)
    print("  Report Generator - Layered KB + Source Tracing")
    print("=" * 50)
    gen = ReportGenerator("./knowledge")
    print(f"  Ready.\n")
    while True:
        try:
            q = input("Topic: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nGoodbye!"); break
        if not q: continue
        if q.lower() in ("exit", "quit", "q"): print("Goodbye!"); break
        try:
            path = gen.generate(q)
            txt = path.replace(".docx", ".txt")
            print(f"  -> {path}")
            print(f"  -> {txt} (preview)")
        except Exception as e:
            import traceback; traceback.print_exc()

if __name__ == "__main__":
    main()
